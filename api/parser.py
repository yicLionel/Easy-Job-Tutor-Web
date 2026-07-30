# -*- coding: utf-8 -*-
"""简历文本抽取：支持 PDF / Word(.docx) / 纯文本(.txt)。"""
import io
import re


def _clean_text(text: str) -> str:
    """清理提取结果，避免全是空白符时被误判为有内容。"""
    return re.sub(r"\n{3,}", "\n\n", (text or "")).strip()


def _extract_pdf_with_pdfplumber(file_bytes: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        return ""
    try:
        parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return _clean_text("\n".join(parts))
    except Exception:
        return ""


def _extract_pdf_with_pypdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        parts = []
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return _clean_text("\n".join(parts))
    except Exception:
        return ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """优先用 pdfplumber，失败后回退到 pypdf。"""
    text = _extract_pdf_with_pdfplumber(file_bytes)
    if text:
        return text
    return _extract_pdf_with_pypdf(file_bytes)


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        import docx
    except ImportError:
        return ""
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return _clean_text("\n".join(parts))
    except Exception:
        return ""


def parse_resume(filename: str, file_bytes: bytes) -> str:
    """解析简历，返回纯文本。若解析失败返回空字符串。"""
    name = (filename or "").lower()
    if name.endswith(".txt"):
        try:
            return _clean_text(file_bytes.decode("utf-8", errors="ignore"))
        except Exception:
            return ""
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if name.endswith((".docx", ".doc")):
        return extract_text_from_docx(file_bytes)
    # 未知类型：两种都试一次
    t = extract_text_from_pdf(file_bytes)
    if not t.strip():
        t = extract_text_from_docx(file_bytes)
    return t
